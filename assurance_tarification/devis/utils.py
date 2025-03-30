from os import truncate
import tempfile
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from .models import Devis

def get_label_garantie(value:str):
    if value=='DO':
        return 'DO seule'
    elif value=='TRC':
        return 'TRC seule'
    else:
        return 'DO + TRC'


def get_label_bien(value:str):
    if value=='Habitation':
        return 'Habitation'
    else:
        return 'Hors Habitation'
        
def get_label_travaux(value:str):
    if value=='Ouvrage_neuf':
        return 'Ouvrage neuf'
    elif value=='Renovation_legere':
        return 'Rénovation légère'
    else:
        'Rénovation lourde'

def generate_word(devis: Devis):
    # Créer un document Word
    doc = Document()
    doc.add_heading('Proposition Commerciale', 0)

    # Ajouter les informations du devis dans le document
    doc.add_paragraph(f"Numéro d'opportunité : {devis.num_opportunite}")
    doc.add_paragraph(f"Nom du client : {devis.nom_client}")
    doc.add_paragraph(f"Tarif proposé : {devis.tarif_propose}€")
    if devis.adresse:
        doc.add_paragraph(f"Adresse : {devis.adresse.numero} {devis.adresse.rue}, {devis.adresse.code_postal} {devis.adresse.ville} ")
    if devis.description:
        doc.add_paragraph(f"Description de l'ouvrage : {devis.description}")
    doc.add_paragraph(f"Type de garantie : {get_label_garantie(devis.type_garantie)}")
    doc.add_paragraph(f"Type de bien : {get_label_bien(devis.type_bien)}")
    doc.add_paragraph(f"Type de travaux : {get_label_travaux(devis.type_travaux)}")
    doc.add_paragraph(f"Déjà existant : {'Oui' if devis.presence else 'Non'}")
    doc.add_paragraph(f"Client VIP : {'Oui' if devis.client_vip else 'Non'}")
    doc.add_paragraph(f"Choix de la RCMO : {'Oui' if devis.rcmo else 'Non'}")

    if(devis.type_garantie == 'TRC' or devis.type_garantie == 'DO_TRC'):
        doc.add_paragraph(f"Taux TRC: {devis.taux_trc}")
    if(devis.type_garantie == 'DO' or devis.type_garantie == 'DO_TRC'):
        doc.add_paragraph(f"Taux DO: {devis.taux_do}")
    
    format_tarif_trc = '{:.2f}'.format(devis.tarif_trc)
    format_tarif_do = '{:.2f}'.format(devis.tarif_do)
    if(devis.type_garantie == 'TRC' or devis.type_garantie == 'DO_TRC'):
        doc.add_paragraph(f"Tarif TRC : {format_tarif_trc} €")
    if(devis.type_garantie == 'DO' or devis.type_garantie == 'DO_TRC'):
        doc.add_paragraph(f"Tarif DO : {format_tarif_do} €")

    format_prime = '{:.2f}'.format(devis.prime_total)
    doc.add_paragraph(f"Prime Total : {format_prime} €" )

    formatted_date_creation = devis.date_creation.strftime("%d-%m-%Y %H:%M:%S")
    doc.add_paragraph(f"Date de création : {formatted_date_creation}")
    
    # Créer un fichier temporaire
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx', mode='wb') as tmp_file:
        file_path = tmp_file.name
        # Sauver le document dans le fichier temporaire
        doc.save(file_path)

        return file_path

def generate_pdf(devis: Devis):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', mode='wb') as tmp_file:
        file_path = tmp_file.name
        c = canvas.Canvas(file_path, pagesize=letter)
        
        c.drawString(100, 750, f"Proposition Commerciale")
        c.drawString(100, 730, f"Numéro d'opportunité : {devis.num_opportunite}")
        c.drawString(100, 710, f"Nom du client : {devis.nom_client}")
        c.drawString(100, 690, f"Tarif proposé : {devis.tarif_propose}€")
        if devis.adresse:
            c.drawString(100, 670, f"Adresse : {devis.adresse.numero} {devis.adresse.rue}, {devis.adresse.code_postal} {devis.adresse.ville} ")
        if devis.description:
            c.drawString(100, 650, f"Description de l'ouvrage : {devis.description}")
        c.drawString(100, 630, f"Type de garantie : {get_label_garantie(devis.type_garantie)}")
        c.drawString(100, 610, f"Type de bien : {get_label_bien(devis.type_bien)}")
        c.drawString(100, 590, f"Type de travaux : {get_label_travaux(devis.type_travaux)}")
        c.drawString(100, 570, f"Déjà existant : {'Oui' if devis.presence else 'Non'}")
        c.drawString(100, 550, f"Client VIP : {'Oui' if devis.client_vip else 'Non'}")
        c.drawString(100, 530, f"Choix de la RCMO : {'Oui' if devis.rcmo else 'Non'}")

        if(devis.type_garantie == 'TRC' or devis.type_garantie == 'DO_TRC'):
            c.drawString(100, 510, f"Taux TRC : {devis.taux_trc}")
        if(devis.type_garantie == 'DO' or devis.type_garantie == 'DO_TRC'):
            c.drawString(100, 490, f"Taux DO : {devis.taux_do}")

        format_tarif_trc = '{:.2f}'.format(devis.tarif_trc)
        format_tarif_do = '{:.2f}'.format(devis.tarif_do)

        if(devis.type_garantie == 'TRC' or devis.type_garantie == 'DO_TRC'):
            c.drawString(100, 470, f"Tarif TRC : {format_tarif_trc} €")
        if(devis.type_garantie == 'DO' or devis.type_garantie == 'DO_TRC'):
            c.drawString(100, 450, f"Tarif DO : {format_tarif_do} €")

        format_prime = '{:.2f}'.format(devis.prime_total)
        c.drawString(100, 430, f"Prime Total : {format_prime} €")

        formatted_date_creation = devis.date_creation.strftime("%d-%m-%Y %H:%M:%S")
        c.drawString(100, 410, f"Date de création : {formatted_date_creation}")

        c.save()
        return file_path